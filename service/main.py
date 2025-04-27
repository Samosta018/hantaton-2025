from gigachat import GigaChat
from dotenv import load_dotenv
import os
import requests
from typing import List, Dict
from docx import Document
import openpyxl
import re
import json
from urllib.parse import quote

load_dotenv()

class GigaCheck:
    def __init__(self):
        self.token = os.getenv("TOKEN_GIGACHAT")
        self.giga = GigaChat(credentials=self.token, verify_ssl_certs=False)
        self.functions = {
            "short_content_gigachat": "Составь краткое описание документа. Ответ в формате: Краткое содержание: *краткое содержимое*. Ничего лишнего!", # Краткое содержание
            "check_spelling": "", # Проверка орфографии
            "check_punctuation_gigachat": "В тексте документа необходимо расставить недостающие запятые и убрать лишние.", # Проверка пунктуации
            "check_create_content_gigachat": "Создай список оглавления для документа (только заголовки, НЕ MARKDOWN). Лишнее не придумывай и не пиши, только СПИСОК.",
            "user_request_gigachat": ""
        }
    
    def extract_text_from_docx(self, path_to_docx): # КОНВЕРТ ИЗ ДОК В СТРОКУ
        doc = Document(path_to_docx)
        full_text = []
        
        for para in doc.paragraphs:
            if not para.text.strip(): 
                continue

            cleaned_text = re.sub(
                r'[a-zA-Z]',
                '', 
                para.text
            )
            cleaned_text = re.sub(
                r'[^а-яА-ЯёЁ0-9\s!?.,:]',
                '', 
                cleaned_text
            )
            
            sentences = re.split(r'(?<=[.!?])\s+', cleaned_text)
            valid_sentences = [s.strip() for s in sentences if s.strip()]
            full_text.extend(valid_sentences)
        
        return ' '.join(full_text)
    
    def check_file(self, path_to_file): # ПРОВЕРКА РАСШИРЕНИЯ ФАЙЛА
        if path_to_file.endswith('.docx'):
            text = self.extract_text_from_docx(path_to_file)
            type_file = 'docx'
        # elif path_to_file.endswith('.xlsx'):
        #     text = extract_text_from_xlsx(path_to_file)
        return text, type_file

    def spell_change_to_docx(self, path_to_file, corrections):
        doc = Document(path_to_file)

        def replace_in_runs(runs, corrections):
            if not runs:
                return

            full_text = ''.join(run.text for run in runs)

            for wrong_word, correction in corrections.items():
                full_text = full_text.replace(wrong_word, correction)

            runs[0].text = full_text
            for run in runs[1:]:
                run.text = ''

        for paragraph in doc.paragraphs:
            replace_in_runs(paragraph.runs, corrections)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_runs(paragraph.runs, corrections)

        for section in doc.sections:
            header = section.header
            footer = section.footer

            for paragraph in header.paragraphs:
                replace_in_runs(paragraph.runs, corrections)
            
            for paragraph in footer.paragraphs:
                replace_in_runs(paragraph.runs, corrections)

        doc.save(path_to_file)

    def spelling_processing(self, text, path_to_file, file_type):
        url = os.getenv("YANDEX_SPELLER_API")
        
        data = {'text': text}
        
        try:
            response = requests.post(url, data=data)
            response.raise_for_status()
            result = response.json()
            
            if not result:
                return "Ошибок не найдено!"
            
            errors = []
            corrections = {}
            
            for item in result:
                wrong_word = item['word']
                if item['s']:
                    correction = item['s'][0]
                    errors.append(f"{wrong_word} → {correction}")
                    corrections[wrong_word] = correction
            
            if not errors:
                return "Ошибок не найдено!"
            
            if file_type == "docx":
                self.spell_change_to_docx(path_to_file, corrections)
            
            return "; ".join(errors)
                
        except Exception as e:
            return {"status": "error", "message": f"Непредвиденная ошибка: {str(e)}"}
    
    def punctuation_change_to_docx(self, path_to_file, word_pairs):
        doc = Document(path_to_file)
        
        required_commas = {tuple(pair) for pair in word_pairs}
        
        def process_runs(runs):
            if not runs:
                return None, 0

            full_text = ''.join(run.text for run in runs)
            original_text = full_text
            modified = False
            changes = []

            comma_pattern = re.compile(r'(\w+)\s*,\s*(\w+)', re.IGNORECASE)
            for match in comma_pattern.finditer(full_text):
                word1, word2 = match.groups()
                if (word1.lower(), word2.lower()) not in required_commas:
                    full_text = full_text[:match.start(1)] + word1 + '  ' + word2 + full_text[match.end(2):]
                    changes.append(f"Удалена лишняя запятая между: {word1} и {word2}")
                    modified = True

            for word1, word2 in word_pairs:
                pattern = re.compile(f'({word1})\s+({word2})', re.IGNORECASE)
                new_text, count = pattern.subn(fr'\1, \2', full_text)
                if count > 0:
                    full_text = new_text
                    changes.append(f"Добавлена запятая между: {word1} и {word2}")
                    modified = True

            if modified:
                runs[0].text = full_text
                for run in runs[1:]:
                    run.text = ''
                return changes
            return None

        all_changes = []
        
        for paragraph in doc.paragraphs:
            changes = process_runs(paragraph.runs)
            if changes:
                all_changes.extend(changes)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        changes = process_runs(paragraph.runs)
                        if changes:
                            all_changes.extend(changes)

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                changes = process_runs(paragraph.runs)
                if changes:
                    all_changes.extend(changes)
            for paragraph in section.footer.paragraphs:
                changes = process_runs(paragraph.runs)
                if changes:
                    all_changes.extend(changes)

        doc.save(path_to_file)
        return all_changes

    def punctuation_processing(self, change_text, path_to_file):
        original_text = self.extract_text_from_docx(path_to_file)
        
        correct_text = change_text
        
        sentences = correct_text.replace(',', ' , ').split('. ')
        word_pairs = []

        for sentence in sentences:
            words = sentence.split()
            for i in range(len(words)-1):
                if words[i] == ',':
                    if i > 0 and i < len(words)-1:
                        pair = [words[i-1], words[i+1]]
                        word_pairs.append(pair)

        changes = self.punctuation_change_to_docx(path_to_file, word_pairs)
        
        if changes:
            report = "Ошибки пунктуации были исправлены!"
        else:
            report = "Все правила пунктуации в документе соблюдены."
        
        return report
    
    
    def delete_file(self, fileid: str) -> None: # УДАЛЕНИЕ ФАЙЛА ИЗ ХРАНИЛИЩА ГИГАЧАТА
        url = f"https://gigachat.devices.sberbank.ru/api/v1/files/{fileid}/delete"
        
        headers = {
            'Accept': 'application/json',
            'Authorization': f'Bearer {self.token}'
        }
        
        response = requests.post(url, headers=headers, data={}, verify=False)
        return response.json()
    
    def analysis_file(self, path_to_file: str, analysis_flags: Dict[str, bool], user_request_text: str = None) -> Dict[str, str]:
        results = {}

        if analysis_flags.get("user_request_gigachat") and user_request_text:
            file = self.giga.upload_file(open(path_to_file, "rb"))
            fileid = file.id_
            result = self.giga.chat({
                "messages": [
                    {
                        "role": "user",
                        "content": user_request_text,
                        "attachments": [fileid],
                    }
                ],
                "temperature": 0.1
            })
            results["user_request_gigachat"] = result.choices[0].message.content
            self.delete_file(fileid)
        
        for analysis_type, should_analyze in analysis_flags.items():
            if not should_analyze:
                continue
                
            if "gigachat" in analysis_type:
                file = self.giga.upload_file(open(path_to_file, "rb"))
                fileid = file.id_
                if "user_request_gigachat" not in analysis_type:
                    result = self.giga.chat({
                        "messages": [
                            {
                                "role": "user",
                                "content": self.functions[analysis_type],
                                "attachments": [fileid],
                            }
                        ],
                        "temperature": 0.1
                    })
                elif "user_request_gigachat" in analysis_type and user_request_text != None:
                    result = self.giga.chat({
                        "messages": [
                            {
                                "role": "user",
                                "content": user_request_text,
                                "attachments": [fileid],
                            }
                        ],
                        "temperature": 0.1
                    })
                    results[analysis_type] = result.choices[0].message.content
                    continue
                
                if "short_content_gigachat" in analysis_type:
                    results[analysis_type] = result.choices[0].message.content
                elif "check_punctuation_gigachat" in analysis_type:
                    results[analysis_type] = self.punctuation_processing(
                        result.choices[0].message.content, 
                        path_to_file
                    )
                elif "check_create_content_gigachat" in analysis_type:
                    results[analysis_type] = result.choices[0].message.content

                self.delete_file(fileid)
                
            elif analysis_type == "check_spelling":
                text, type_file = self.check_file(path_to_file)
                results[analysis_type] = self.spelling_processing(text, path_to_file, type_file)

        return results
    

gigacheck = GigaCheck()
